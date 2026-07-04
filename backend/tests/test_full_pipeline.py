from __future__ import annotations

from pathlib import Path

from src.mekg.models import CrossCorpusSearchRequest, ElementKind
from src.mekg.parsers import DocumentParser
from src.mekg.pipeline import FullCorpusPipeline
from src.mekg.retrieval import CrossCorpusRetriever
from src.mekg.search_store import corpus_for_category, vector_literal

from .test_mekg import DisabledVision, config


class FastPage:
    def get_text(self, mode):
        assert mode == "text"
        return "Fast text layer with nickel recovery 94 percent at 25 C. " * 5

    def get_images(self, **_kwargs):
        raise AssertionError("fast parser must not inspect PDF images")

    def find_tables(self):
        raise AssertionError("fast parser must not call find_tables")


class FastPdf:
    def __iter__(self):
        return iter([FastPage()])


def test_fast_pdf_uses_only_text_layer(tmp_path, monkeypatch):
    source = tmp_path / "fast.pdf"
    source.write_bytes(b"fake-pdf-for-hash")
    monkeypatch.setattr("src.mekg.parsers.fitz.open", lambda _path: FastPdf())
    parser = DocumentParser(config(tmp_path), DisabledVision())

    parsed = parser.parse(source, source_locator="Доклады/fast.pdf", category="Доклады", fast=True)

    assert parsed.elements
    assert all(item.kind == ElementKind.TEXT for item in parsed.elements)
    assert not list((tmp_path / "artifacts" / parsed.version_id).rglob("figures"))


def test_openxml_word_fallback_keeps_text_and_native_tables(tmp_path):
    from docx import Document
    from src.mekg.models import ParsedDocument

    source = tmp_path / "fallback.docm"
    word = Document()
    word.add_paragraph("Nickel recovery was 94 percent at 25 C.")
    table = word.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "recovery"
    table.cell(0, 1).text = "94 %"
    word.save(source)
    parser = DocumentParser(config(tmp_path), DisabledVision())
    document = ParsedDocument(
        document_id="d", version_id="v", source_locator="fallback.docm",
        file_name=source.name, file_type="docm", sha256="x", size_bytes=source.stat().st_size,
    )

    parser._parse_openxml_word(source, document)

    assert any(item.kind == ElementKind.TEXT and "Nickel recovery" in item.text for item in document.elements)
    assert any(item.kind == ElementKind.TABLE_ROW and "94 %" in item.text for item in document.elements)


def test_all_five_corpora_have_stable_mapping():
    assert corpus_for_category("Доклады") == "internal_reports"
    assert corpus_for_category("Журналы") == "scientific_journals"
    assert corpus_for_category("Материалы конференций") == "conference_materials"
    assert corpus_for_category("Обзоры") == "reviews"
    assert corpus_for_category("Статьи") == "scientific_articles"


def test_pgvector_literal_preserves_768_dimensions():
    encoded = vector_literal([0.25] * 768)
    assert encoded.startswith("[") and encoded.endswith("]")
    assert encoded.count(",") == 767


def test_rrf_and_diversity_limit_documents_and_corpora():
    request = CrossCorpusSearchRequest(query="nickel recovery", final_k=10)
    rows = [
        {
            "chunk_id": f"c{index}", "document_id": "same" if index < 5 else f"d{index}",
            "corpus_id": "scientific_articles", "text": "nickel", "dense_score": 1 - index / 100,
        }
        for index in range(12)
    ]
    fused = CrossCorpusRetriever._fuse([("scientific_articles", "dense", rows)], request)
    diversified = CrossCorpusRetriever._diversify(fused, 10, "research")
    assert len([item for item in diversified if item["document_id"] == "same"]) == 3
    assert len(diversified) <= 8


def test_numeric_strict_matching_and_slots():
    wanted = CrossCorpusRetriever._numeric_mentions("recovery 94 % at 25 C")
    assert CrossCorpusRetriever._numeric_score(wanted, "Measured recovery was 94 %.") == 1
    assert CrossCorpusRetriever._matched_slots(["nickel recovery", "pressure"], "Nickel recovery was measured") == ["nickel recovery"]


def test_evidence_selection_is_bounded_to_text_and_native_rows():
    from src.mekg.models import ParsedDocument, SourceElement

    elements = [
        SourceElement(id=f"c{i}", kind=ElementKind.TEXT, text=("Nickel recovery 94 percent. " * 3) + str(i))
        for i in range(8)
    ]
    elements.append(SourceElement(id="figure", kind=ElementKind.FIGURE, text="A figure description " * 10))
    document = ParsedDocument(
        document_id="doc", version_id="v", source_locator="x", file_name="x.pdf",
        file_type="pdf", sha256="0" * 64, size_bytes=1, elements=elements,
    )
    selected = FullCorpusPipeline._evidence_elements(document)[:6]
    assert len(selected) == 6
    assert all(item.kind in {ElementKind.TEXT, ElementKind.TABLE_ROW} for item in selected)
